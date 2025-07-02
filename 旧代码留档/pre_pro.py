import os
import re
import sys
import shutil
import csv
import hashlib
import pdfplumber
import fitz
import io
import pandas as pd
import tempfile
import logging
import xml.etree.ElementTree as ET
from PIL import Image
from docx import Document

logging.getLogger('pdfminer').setLevel(logging.ERROR)
os.environ['UNRAR_LIB_PATH'] = os.path.expanduser('~/.local/lib/libunrar.so')
# 设置UTF-8环境变量
os.environ['LANG'] = 'zh_CN.UTF-8'
os.environ["PYMUPDF_SILENT"] = "yes" 

from unrar import rarfile
import zipfile
import tarfile


MIN_SID_LENGTH = 10
SUPPORTED_EXTENSIONS = ['.zip', '.rar', '.tar', '.gz', '.bz2']
KEIL_PROJECT_EXTS = ('.uvproj', '.uvprojx')
CODE_EXTENSIONS = ('.c', '.h')


# def extract_student_info(filename):

#     # if '组-' not in filename:
#     #     return None, []
    
#     # 获取组信息和获取“组-”后面的学生姓名以及学号信息
#     group_match = re.search(r'第(\d{1,3})组[_-]?(.*)', filename)
#     if not group_match:
#         return None, []

#     group_number, student_section = group_match.groups()

#     # 获取“组-”后面的学生姓名以及学号信息
#     # try:
#     #     student_section = filename.split('组-')[-1]
#     # except IndexError:
#     #     return []

#     student_pattern = re.compile(
#         r'(?<!\d)'  # 确保学号前不紧接数字
#         r'(\d{10})'  # 匹配10位学号
#         r'[\s_、，-]*'  # 允许常见分隔符
#         r'([\u4e00-\u9fa5]{2,})'  # 匹配中文姓名
#         r'|'  # 或反向顺序
#         r'([\u4e00-\u9fa5]{2,})'  # 匹配中文姓名
#         r'[\s_、，-]*' 
#         r'(\d{10})'  # 匹配10位学号
#         r'(?!\d)'  # 确保学号后不紧接数字
#     )


#     cleaned_section = re.sub(r'[、，]', ' ', student_section)
#     # 查找所有匹配项
#     matches = student_pattern.finditer(cleaned_section)
    
#     # # 获取所有符合条件的内容
#     # matches = student_pattern.findall(student_section)


#     valid_students = []
#     for match in matches:
#         if match.group(1) and match.group(2):  # 学号+姓名
#             sid, name = match.group(1), match.group(2)
#         elif match.group(3) and match.group(4):  # 姓名+学号
#             name, sid = match.group(3), match.group(4)
#         else:
#             continue  # 无效匹配
        
#         # 验证学号和姓名
#         if len(sid) == 10 and 2 <= len(name) <= 4:  # 假设姓名长度2-4字
#             valid_students.append((sid, name))
#         else:
#             print(f"无效条目: {match.group()} (学号长度:{len(sid)}, 姓名长度:{len(name)})")

#     # for sid, name in matches:
#     #     # 确认正常内容
#     #     if len(name) >= 2 and len(sid) >= MIN_SID_LENGTH:
#     #         valid_students.append((sid, name))
    
#     # return f"第{group_number}组", valid_students
#     # valid_students = []
#     # for entry in final_entries:
#     #     match = student_pattern.fullmatch(entry)
#     #     if not match:
#     #         print(f"格式错误: {entry}")
#     #         continue
            
#     #     sid, name = (match.group(1), match.group(2)) if match.group(1) else (match.group(4), match.group(3))
        
#     #     if len(sid) >= MIN_SID_LENGTH and len(name) >= 2:
#     #         valid_students.append((sid, name))
#     #     else:
#     #         print(f"无效条目: {entry}")
    
#     return f"第{group_number}组", valid_students

# 提取报告函数
def extract_file(dir_path, group_output_dir):
    # group_name = os.path.basename(dir_path)
    # report_files = []
    report_output_dir = os.path.join(group_output_dir, "实验报告")
    os.makedirs(report_output_dir, exist_ok=True)


    group_files = {}
    for root, dirs, files in os.walk(dir_path):
        dirs[:] = [d for d in dirs if not d.startswith(('.', '_'))]
        group_name = os.path.basename(root)
        for file in files:
            if (file.lower().endswith(('.pdf', '.doc', '.docx'))) and \
               ("代码" not in file and "源码" not in file):
                file_path = os.path.join(root, file)
                if group_name not in group_files:
                    group_files[group_name] = []
                # report_files.append((group_name,))
                group_files[group_name].append(file_path)
        
    
    for group_name, report_files in group_files.items():
        if len(report_files) == 1:
            report = report_files[0]
        else:
            sorted_files = sorted(report_files, key=lambda f: (
                0 if f.lower().endswith('.pdf') else
                1 if f.lower().endswith('.docx') else
                2
            ))
            report = sorted_files[0]
        extract_format(group_name, report, report_output_dir)

def extract_format(group_name, file_path, output_dir):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        extract_pdf_content(file_path, output_dir)
    elif ext == '.docx':
        extract_docx_content(file_path, output_dir)
    elif ext == '.doc':
        print("doc格式文件效果有限，建议转化为pdf或者docx格式")
        extract_doc_content(file_path, output_dir)
    else:
        # raise ValueError(f"报告非doc或pdf格式： {ext}")
        print(f"\n无法提取改组实验报告：{group_name}\n")

def extract_pdf_content(pdf_path, output_dir):

    # os.makedirs(output_dir, exist_ok=True)
    # group_path = os.path.join(output_dir, group_name)
    # os.makedirs(group_path, exist_ok=True)
    text_dir = os.path.join(output_dir, "文字")
    img_dir = os.path.join(output_dir, "图片")
    table_dir = os.path.join(output_dir, "表格")
    os.makedirs(text_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(table_dir, exist_ok=True)



    doc = fitz.open(pdf_path)
    all_text = []
    image_count = 0

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        all_text.append(text)

        # with open(os.path.join(text_dir, f"page_{page_num}.txt"), "w", encoding="utf-8") as f:
        #     f.write(text)

    full_text = "\n".join(all_text)
    with open(os.path.join(text_dir, "完整文本.txt"), "w", encoding="utf-8") as f:
        f.write(full_text)

    extract_table_content(pdf_path, table_dir)



    for page_num in range (len(doc)):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)

        for image_inedx, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            img_ext = base_image["ext"]
            img_path = os.path.join(img_dir, f"page_{page_num+1}_img_{image_inedx+1}.{img_ext}")
            with open(img_path, "wb") as f:
                f.write(image_bytes)

            image_count += 1

def extract_table_content(pdf_path, table_dir):
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            table = page.extract_tables()

            if table:
                # print(1)
                for table_num, table in enumerate(table):
                    cleaned_table = []
                    for row in table:
                        cleaned_row = [cell.replace("\n", " ") if cell else "" for cell in row]
                        cleaned_table.append(cleaned_row)
                    
                    csv_path = os.path.join(table_dir, f"page_{page_num+1}_table_{table_num+1}")

                    with open(csv_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.writer(f)
                        writer.writerows(cleaned_table)

def extract_docx_images(temp_dir, img_dir):
    """从解压的DOCX文件中提取图片"""
    media_path = os.path.join(temp_dir, 'word', 'media')
    if os.path.exists(media_path):
        img_count = 1
        for img_file in os.listdir(media_path):
            src_path = os.path.join(media_path, img_file)
            if os.path.isfile(src_path):
                # 获取文件扩展名
                _, ext = os.path.splitext(img_file)
                dest_path = os.path.join(img_dir, f"image_{img_count}{ext}")
                shutil.copy(src_path, dest_path)
                img_count += 1

def extract_docx_tables(docx_path, table_dir):
    try:
        doc = Document(docx_path)
        table_count = 1
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs)
                    row_data.append(cell_text.strip().replace('\n', ' '))
                table_data.append(row_data)
            

            if table_data and any(any(cell for cell in row) for row in table_data):
                csv_path = os.path.join(table_dir, f"table_{table_count}.csv")
                with open(csv_path, "w", encoding="utf-8", newline="") as f:
                    writer = csv.writer(f)
                    writer.writerows(table_data)
                table_count += 1
    except ImportError:
        print("警告：未安装python-docx库，跳过表格提取")
    except Exception as e:
        print(f"提取表格时出错: {e}")

def extract_xml_text(xml_path, namespaces):
    """从XML文件中提取文本内容"""
    try:
        tree = ET.parse(xml_path)
        root = tree.getroot()
        
        # 提取所有文本内容
        text_parts = []
        for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                text_parts.append(t.text.strip())
        
        return " ".join(text_parts)
    except Exception as e:
        print(f"解析XML文件 {xml_path} 时出错: {e}")
        return ""


def extract_docx_text(temp_dir):
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }
    text_parts = []

    document_path = os.path.join(temp_dir, 'word', 'document.xml')
    if os.path.exists(document_path):
        text_parts.append(extract_xml_text(document_path, ns))

    header_dir = os.path.join(temp_dir, 'word', 'header')
    if os.path.exists(header_dir):
        for header_file in os.listdir(header_dir):
            if header_file.endswith('.xml'):
                header_path = os.path.join(header_dir, header_file)
                text_parts.append(extract_xml_text(header_path, ns))

    footer_dir = os.path.join(temp_dir, 'word', 'footer')
    if os.path.exists(footer_dir):
        for footer_file in os.listdir(footer_dir):
            if footer_file.endswith('.xml'):
                footer_path = os.path.join(footer_dir, footer_file)
                text_parts.append(extract_xml_text(footer_path, ns))
    
    return "\n\n".join(text_parts)


def extract_docx_content(docx_path, output_dir):

    # os.makedirs(output_dir, exist_ok=True)
    # group_path = os.path.join(output_dir, group_name)
    # os.makedirs(group_path, exist_ok=True)
    text_dir = os.path.join(output_dir, "文字")
    img_dir = os.path.join(output_dir, "图片")
    table_dir = os.path.join(output_dir, "表格")
    os.makedirs(text_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(table_dir, exist_ok=True)

    tmp_dir = tempfile.mkdtemp()

    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            zip_file.extractall(tmp_dir)

        text_content = extract_docx_text(tmp_dir)
        with open(os.path.join(text_dir, "完整文本.txt"), "w", encoding="utf-8") as f:
            f.write(text_content)

        extract_docx_images(tmp_dir, img_dir)

        extract_docx_tables(docx_path, table_dir)
    
    except Exception as e:
        print(f"处理docx文件时出错： {e}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def extract_doc_content(doc_path, output_dir):
    """提取DOC文件内容（基本文本提取）"""
    # os.makedirs(output_dir, exist_ok=True)
    text_dir = os.path.join(output_dir, "文字")
    img_dir = os.path.join(output_dir, "图片")
    table_dir = os.path.join(output_dir, "表格")
    os.makedirs(text_dir, exist_ok=True)
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(table_dir, exist_ok=True)
    
    # 尝试使用antiword提取文本（Linux）
    if sys.platform == 'linux' or sys.platform == 'linux2':
        try:
            from subprocess import check_output
            text = check_output(['antiword', doc_path]).decode('utf-8', errors='ignore')
            with open(os.path.join(text_dir, "完整文本.txt"), "w", encoding="utf-8") as f:
                f.write(text)
            return
        except Exception:
            pass
    
    # 基本替代方案
    with open(os.path.join(text_dir, "完整文本.txt"), "w", encoding="utf-8") as f:
        f.write("DOC格式提取需要安装antiword（Linux）或textract（Windows/Mac）")
    
    with open(os.path.join(img_dir, "说明.txt"), "w", encoding="utf-8") as f:
        f.write("DOC格式不支持图片提取")
    
    with open(os.path.join(table_dir, "说明.txt"), "w", encoding="utf-8") as f:
        f.write("DOC格式不支持表格提取")

def extract_table_content(pdf_path, table_dir):
    """提取PDF中的表格"""
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            
            if tables:
                for table_num, table in enumerate(tables):
                    cleaned_table = []
                    for row in table:
                        cleaned_row = [cell.replace("\n", " ") if cell else "" for cell in row]
                        cleaned_table.append(cleaned_row)
                    
                    csv_path = os.path.join(table_dir, f"page_{page_num+1}_table_{table_num+1}.csv")
                    
                    with open(csv_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.writer(f)
                        writer.writerows(cleaned_table)



# 组信息
def extract_info(folder_name):

    student_ids = re.findall(r'\d{10}', folder_name)
    if not student_ids:
        return None, None 
    
    names = re.findall(r'[\u4e00-\u9fa5]{2,4}', folder_name)
    
    students = []
    

    if len(student_ids) == len(names):

        students = list(zip(student_ids, names))
    else:

        for sid in student_ids:

            sid_pos = folder_name.find(sid)
            
            closest_name = None
            min_distance = float('inf')
            
            for name in names:
                name_pos = folder_name.find(name)

                distance = abs(name_pos - sid_pos)
                if distance < min_distance:
                    min_distance = distance
                    closest_name = name
            
            if closest_name:
                students.append((sid, closest_name))
                names.remove(closest_name)  
    
    cleaned_info = re.sub(
        r'((?!\d{10}|[\u4e00-\u9fa5]{2,4}|和|与|及).)+', 
        ' ', 
        folder_name
    )
    
    return students, cleaned_info.strip()



def find_keil_project(root_dir):
    for root, dirs, files in os.walk(root_dir):
        # dirs[:] = [d for d in dirs if not d.startswith('EXTRACTED_')]
        for file in files:
            if file.lower().endswith(KEIL_PROJECT_EXTS):
                return os.path.join(root, file)
    return None


def safe_unzip(file_path, extract_dir):

    try: 
        os.makedirs(extract_dir, exist_ok=True)

        if file_path.endswith('.zip'):
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
        elif file_path.endswith('.rar'):
            with rarfile.RarFile(file_path, 'r') as rar_ref:
                rar_ref.extractall(extract_dir)

            
        elif any(file_path.endswith(ext) for ext in ['.tar', '.tar.gz', '.tar.bz2']):
            with tarfile.open(file_path, 'r:*', encoding='gbk') as tar_ref:
                tar_ref.extractall(extract_dir)

        
        print(f"解压文件: {os.path.basename(file_path)}")
        os.remove(file_path)
        return True

    except Exception as e:
        print(f"解压失败 {os.path.basename(file_path)}: {str(e)}")
        return False



def unzip_files(entry_path):

    # 确认所有压缩文件都已经被解压
    flag = False

    for root, dirs, files in os.walk(entry_path, topdown=True):
        dirs[:] = [d for d in dirs if not d.startswith('EXTRACTED_')]

        for file in files:
            if any(file.lower().endswith(ext) for ext in SUPPORTED_EXTENSIONS):
                file_path = os.path.join(root, file)
                extract_dir = os.path.join(
                    root,
                    f"EXTRACTED_{os.path.splitext(file)[0]}"
                )
                
                if safe_unzip(file_path, extract_dir):
                    flag = True

    if flag:
        unzip_files(entry_path)

# 新增功能：提取PDF中的代码
def extract_code_from_pdf(pdf_path):
    """使用pdfplumber提取PDF中的C代码"""
    try:
        full_text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
        
        # 清理代码：移除页码等无关内容
        cleaned_text = re.sub(r'第\s*\d+\s*[页頁]', '', full_text)
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
        
        return cleaned_text
    except Exception as e:
        print(f"提取PDF代码失败: {pdf_path}, error: {str(e)}")
        return ""

# 新增功能：检测并处理包含"代码"的PDF文件
def process_pdf_files(root_dir):
    """查找并处理包含'代码'的PDF文件"""
    pdf_files = []
    for root, _, files in os.walk(root_dir):
        for file in files:
            if (file.lower().endswith('.pdf') and "代码" in file) or (file.lower().endswith('.pdf') and "源码" in file):
                pdf_files.append(os.path.join(root, file))
    
    generated_c_files = []
    for pdf_file in pdf_files:
        print(f"发现代码PDF文件: {os.path.basename(pdf_file)}")
        code_text = extract_code_from_pdf(pdf_file)
        
        if code_text.strip():
            # 生成对应的.c文件名
            c_filename = os.path.splitext(os.path.basename(pdf_file))[0] + ".c"
            c_filepath = os.path.join(os.path.dirname(pdf_file), c_filename)
            
            # 保存提取的代码
            with open(c_filepath, 'w', encoding='utf-8') as f:
                f.write(code_text)
            print(f"  已提取代码保存至: {c_filename}")
            generated_c_files.append(c_filepath)
    
    return generated_c_files

def collect_Ccode(code_root):
    code_files = []
    for root, _, files in os.walk(code_root):
        for file in files:
            if file.lower().endswith(CODE_EXTENSIONS):
                code_files.append(os.path.join(root, file))
    return code_files



# # def integrate_c_files(group_number, students_info, code_files, output_base):
# def integrate_c_files(students_info, code_files, output_base):
#     # member_list = '_'.join([f"{name}_{sid}" for sid, name in students_info])
#     # group_tag = f"{group_number}_{member_list}"
    
#     member_list = '_'.join([f"{name}_{sid}" for sid, name in students_info])
#     output_dir = os.path.join(output_base, member_list)


def process_submitted_files(entry_path, output_base):
    # 如果上交的存在压缩包则先解压压缩包
    # try:
    #     print(f"\n处理提交数据，提取代码：{os.path.basename(entry_path)}")
    #     unzip_files(entry_path)
    #     return True
    # except Exception as e:
    #     print(f"处理失败： {os.path.basename(entry_path)} - {str(e)}")
    #     return False
    print(f"\n处理提交数据，提取代码：{os.path.basename(entry_path)}")
    unzip_files(entry_path)

    # print(entry_path)
    # input()


    # group_number, students_info = extract_student_info(os.path.basename(entry_path))
    students_info, _ = extract_info(entry_path)
    if not students_info:
        print(f"警告，缺少有效姓名或者学号：{os.path.basename(entry_path)}")
        return False
    # print(students_info)
    member_list = '_'.join([f"{name}_{sid}" for sid, name in students_info])
    # group_dir_name = f"{group_number}_{member_list}"
    group_output_dir = os.path.join(output_base, member_list)
    os.makedirs(group_output_dir, exist_ok=True)
    # print(group_output_dir)
    # print(group_dir_name)
    # exit(0)

    keil_project = find_keil_project(entry_path)    
    code_root = os.path.dirname(keil_project) if keil_project else entry_path
    # print(keil_project)
    # if keil_project:
        # print(f"找到Keil项目文件： {os.path.basename(keil_project)}")
        # print(f"找到Keil项目文件： {keil_project}")
        # code_root = os.path.dirname(keil_project)
    # else:
        # print(f"未找到Keil项目文件")
        # code_root = entry_path

    code_files = collect_Ccode(code_root)
    # print(code_files)

    if not code_files:
        generated_c_files = process_pdf_files(entry_path)
        for c_file in generated_c_files:
            if c_file not in code_files:
                code_files.append(c_file)


    if code_files:
        output_code_dir = os.path.join(group_output_dir, '代码')
        os.makedirs(output_code_dir, exist_ok=True)

        processed_files = []
        for src_file in code_files:
            try:
                # 创建目标路径
                base_name = os.path.basename(src_file)
                # dest_name = f"{member_list}_{base_name}"
                dest_name = f"{base_name}"
                dest_path = os.path.join(output_code_dir, dest_name)
                
                # 处理重名文件
                counter = 1
                while os.path.exists(dest_path):
                    # 保留原始文件名前缀，只在必要时添加序号
                    dest_name = f"{os.path.splitext(base_name)[0]}_{counter}{os.path.splitext(base_name)[1]}"
                    dest_path = os.path.join(output_code_dir, dest_name)
                    counter += 1
                
                # 复制文件并保留元数据
                shutil.copy2(src_file, dest_path)
                processed_files.append({
                    "source": src_file,
                    "destination": dest_path,
                    "size": os.path.getsize(src_file),
                    "md5": hashlib.md5(open(src_file, 'rb').read()).hexdigest()
                })
                
            except Exception as e:
                print(f"文件整合失败: {src_file} -> {str(e)}")
                # continue

        # 记录元数据
        metadata_path = os.path.join(group_output_dir, "metadata.csv")
        
        with open(metadata_path, 'a', encoding='utf-8') as f:
            # writer = csv.DictWriter(f, fieldnames=["group", "members", "src_path", "dest_path", "size", "md5"])
            writer = csv.DictWriter(f, fieldnames=["成员", "源文件路径", "目标文件路径", "文件大小", "MD5"])
            writer.writeheader()
            
            for record in processed_files:
                writer.writerow({
                    # "group": group_number,
                    "成员": member_list,
                    "源文件路径": record["source"],
                    "目标文件路径": record["destination"],
                    "文件大小": record["size"],
                    "MD5": record["md5"]
                })

    else: 
        print(f"无有效代码文件或使用文档等其他文件格式")
        return False
    

    # group_output = integrate_c_files(group_number, students_info, code_files, output_dir)
    # group_output = integrate_c_files(students_info, code_files, output_dir)
    # integrate_c_files(students_info, code_files, output_dir)

    # report_output_dir = os.path.join(output_dir, "实验报告")
    print(f"开始提取实验报告：{os.path.basename(entry_path)}")
    try:
        extract_file(entry_path, group_output_dir)
    except Exception as e:
        print(f"报告提取失败：{str(e)}")
    return True

    return group_output_dir



def pre_process_files(source_dir, output_dir="test_total_files"):
    os.makedirs(output_dir, exist_ok=True)


    for entry in os.listdir(source_dir):
        entry_dir = os.path.join(source_dir, entry)

        if not os.path.isdir(entry_dir):
            print(f"跳过非目录项: {entry}")
            continue
        process_submitted_files(entry_dir, output_dir)




if __name__ == "__main__":
    if len(sys.argv) > 1:
        soc_dir = sys.argv[1]

    if not os.path.exists(soc_dir):
        print(f"Error, no such dictory")
        sys.exit(1)
    
    print("Preprocessing Submissions")
    print(f"source dictory: {soc_dir}")


    # print(sys.version_info)
    pre_process_files(soc_dir)